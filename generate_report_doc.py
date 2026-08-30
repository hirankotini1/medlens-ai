import os 
from docx import Document 
from docx .shared import Inches ,Pt ,RGBColor 
from docx .enum .text import WD_ALIGN_PARAGRAPH 
from docx .enum .table import WD_TABLE_ALIGNMENT 
from docx .oxml import OxmlElement ,parse_xml 
from docx .oxml .ns import nsdecls ,qn 

def set_cell_background (cell ,fill_color ):
    tcPr =cell ._tc .get_or_add_tcPr ()
    shd =parse_xml (f'<w:shd {nsdecls ("w")} w:fill="{fill_color }"/>')
    tcPr .append (shd )

def create_report ():
    doc =Document ()


    sections =doc .sections 
    for section in sections :
        section .top_margin =Inches (1 )
        section .bottom_margin =Inches (1 )
        section .left_margin =Inches (1 )
        section .right_margin =Inches (1 )


    normal_style =doc .styles ['Normal']
    normal_style .font .name ='Calibri'
    normal_style .font .size =Pt (11 )
    normal_style .font .color .rgb =RGBColor (30 ,41 ,59 )
    normal_style .paragraph_format .line_spacing =1.15 
    normal_style .paragraph_format .space_after =Pt (6 )

    def add_title (text ,subtitle =None ):
        p =doc .add_paragraph ()
        p .alignment =WD_ALIGN_PARAGRAPH .CENTER 
        run =p .add_run (text )
        run .font .size =Pt (26 )
        run .font .bold =True 
        run .font .color .rgb =RGBColor (79 ,70 ,229 )
        if subtitle :
            p2 =doc .add_paragraph ()
            p2 .alignment =WD_ALIGN_PARAGRAPH .CENTER 
            run2 =p2 .add_run (subtitle )
            run2 .font .size =Pt (14 )
            run2 .font .color .rgb =RGBColor (100 ,116 ,139 )
            p2 .paragraph_format .space_after =Pt (24 )

    def add_h1 (text ):
        p =doc .add_paragraph ()
        p .paragraph_format .space_before =Pt (18 )
        p .paragraph_format .space_after =Pt (6 )
        p .paragraph_format .keep_with_next =True 
        run =p .add_run (text )
        run .font .size =Pt (18 )
        run .font .bold =True 
        run .font .color .rgb =RGBColor (30 ,41 ,59 )

    def add_h2 (text ):
        p =doc .add_paragraph ()
        p .paragraph_format .space_before =Pt (12 )
        p .paragraph_format .space_after =Pt (4 )
        p .paragraph_format .keep_with_next =True 
        run =p .add_run (text )
        run .font .size =Pt (13.5 )
        run .font .bold =True 
        run .font .color .rgb =RGBColor (79 ,70 ,229 )

    def add_callout (text ,title ="NOTE"):
        tbl =doc .add_table (rows =1 ,cols =1 )
        tbl .alignment =WD_TABLE_ALIGNMENT .CENTER 
        cell =tbl .cell (0 ,0 )
        set_cell_background (cell ,"F1F5F9")
        p =cell .paragraphs [0 ]
        p .paragraph_format .space_before =Pt (4 )
        p .paragraph_format .space_after =Pt (4 )
        r_t =p .add_run (f"[{title }] ")
        r_t .font .bold =True 
        r_t .font .color .rgb =RGBColor (99 ,102 ,241 )
        r_b =p .add_run (text )
        r_b .font .size =Pt (10 )
        doc .add_paragraph ().paragraph_format .space_after =Pt (4 )




    p_pre =doc .add_paragraph ()
    p_pre .paragraph_format .space_before =Pt (72 )

    add_title ("NEXUS PATHOLOGY","A Digital Pathology Laboratory Management Platform with Experimental Machine-Learning Decision Support")

    p_meta =doc .add_paragraph ()
    p_meta .alignment =WD_ALIGN_PARAGRAPH .CENTER 
    p_meta .add_run ("A Project Report Submitted in Partial Fulfillment of the Requirements\nfor the Degree of\n\n").font .size =Pt (11 )
    r_deg =p_meta .add_run ("BACHELOR OF TECHNOLOGY\nIN\nCOMPUTER SCIENCE AND ENGINEERING\n\n")
    r_deg .font .bold =True 
    r_deg .font .size =Pt (12 )

    p_sub =doc .add_paragraph ()
    p_sub .alignment =WD_ALIGN_PARAGRAPH .CENTER 
    p_sub .paragraph_format .space_before =Pt (36 )
    p_sub .add_run ("Submitted By:\n").font .bold =True 
    p_sub .add_run ("[Student Name 1] (Roll No: [XXXX])\n[Student Name 2] (Roll No: [XXXX])\n[Student Name 3] (Roll No: [XXXX])\n\n")
    p_sub .add_run ("Under the Guidance of:\n").font .bold =True 
    p_sub .add_run ("[Faculty Guide Name / Designation]\n\n")
    p_sub .add_run ("DEPARTMENT OF COMPUTER SCIENCE & ENGINEERING\n[COLLEGE / UNIVERSITY NAME]\n[ACADEMIC YEAR 2025–2026]").font .bold =True 

    doc .add_page_break ()




    add_h1 ("Certificate")
    p_cert =doc .add_paragraph ()
    p_cert .paragraph_format .space_before =Pt (12 )
    p_cert .add_run (
    "This is to certify that the project report entitled 'NEXUS PATHOLOGY: A Digital Pathology Laboratory Management Platform with Experimental Machine-Learning Decision Support' is a bonafide record of work carried out by [Student Name(s)] under my supervision and guidance, in partial fulfillment of the requirements for the award of the Degree of Bachelor of Technology in Computer Science and Engineering.\n\n"
    "To the best of my knowledge, the matter embodied in this report has not been submitted to any other University or Institute for the award of any degree or diploma."
    )
    p_sig =doc .add_paragraph ()
    p_sig .paragraph_format .space_before =Pt (72 )
    p_sig .add_run (
    "_________________________\t\t\t\t_________________________\n"
    "[Internal Guide Name]\t\t\t\t[Head of Department]\n"
    "Department of CSE\t\t\t\tDepartment of CSE"
    )
    doc .add_page_break ()




    add_h1 ("Declaration")
    doc .add_paragraph (
    "We hereby declare that the project entitled 'NEXUS PATHOLOGY: A Digital Pathology Laboratory Management Platform with Experimental Machine-Learning Decision Support' submitted to [College / University Name] is a record of original work done by us under the supervision of [Guide Name].\n\n"
    "We further declare that this report has not previously formed the basis for the award of any degree, diploma, associateship, fellowship, or other similar title to the best of our knowledge."
    )
    p_dec_sig =doc .add_paragraph ()
    p_dec_sig .paragraph_format .space_before =Pt (72 )
    p_dec_sig .add_run (
    "Date: __________________\n"
    "Place: _________________\t\t\t\t[Signature of Candidate(s)]"
    )
    doc .add_page_break ()




    add_h1 ("Acknowledgement")
    doc .add_paragraph (
    "We express our deep sense of gratitude to our respected Principal, Head of Department, and project guide [Guide Name] for their invaluable guidance, constant encouragement, and constructive criticism throughout the development of this project.\n\n"
    "We would also like to thank the faculty and laboratory staff of the Department of Computer Science and Engineering for providing the computational resources necessary to design, train, evaluate, and test our machine learning pipelines and web architecture.\n\n"
    "Finally, we extend our heartfelt gratitude to our families and peers for their continuous support and understanding."
    )
    doc .add_page_break ()




    add_h1 ("Abstract")
    doc .add_paragraph (
    "Clinical pathology laboratories generate diagnostic data essential for medical evaluations. Traditional workflows often rely on physical paper records or siloed electronic systems with limited patient accessibility and no integrated computational decision support. Furthermore, applying machine learning to clinical workflows introduces challenges regarding data privacy, report integrity, and diagnostic overreach.\n\n"
    "Nexus Pathology is a web-based digital pathology laboratory management platform integrated with an experimental machine-learning (ML) decision-support backend. The platform provides a secure, role-based environment with two primary interfaces: an Administrative Portal for laboratory staff to manage patients, record structured laboratory parameters, and control report lifecycles; and an authenticated Patient Portal providing privacy-preserving access to official laboratory reports with reference ranges and doctor remarks.\n\n"
    "The system incorporates five independently trained and validated machine-learning pipelines: Complete Blood Count / Anemia (Logistic Regression), Dengue Hematology (Random Forest), Liver Function (Gradient Boosting), Thyroid Profile (Multinomial Logistic Regression), and Malaria Microscopy Smears (OpenCV 354-dimensional feature extraction + Gradient Boosting). A core architectural pillar is the strict decoupling of official laboratory records from ML predictions, ensuring medical reports remain immutable while persisting ML inferences to an independent audit log. The system is verified with 25 automated security, integration, and ML test scenarios achieving a 100% pass rate."
    )
    add_callout ("Educational & Research Boundary: Machine-learning outputs in Nexus Pathology serve as decision support and do not constitute an autonomous medical diagnosis.","ETHICAL NOTICE")
    doc .add_page_break ()




    add_h1 ("Table of Contents")
    toc_items =[
    ("1. Introduction","1"),
    ("2. Problem Statement","2"),
    ("3. Project Objectives","3"),
    ("4. Existing vs. Proposed System","4"),
    ("5. System Requirements Specification","5"),
    ("6. System Architecture","6"),
    ("7. Database Design & Schema","8"),
    ("8. Implementation Details","10"),
    ("9. Machine Learning Methodology","12"),
    ("10. Dataset Description & Features","14"),
    ("11. Model Training & Validation","16"),
    ("12. Model Performance & Evaluation","18"),
    ("13. Controlled Synthetic Data Experiment","20"),
    ("14. Cybersecurity & Data Privacy","22"),
    ("15. System Testing & Quality Assurance","24"),
    ("16. Results & Discussion","26"),
    ("17. System Limitations","28"),
    ("18. Future Scope","29"),
    ("19. Conclusion","30"),
    ("20. References","31"),
    ("21. Appendix & Viva Questions","32")
    ]
    for title ,page in toc_items :
        p =doc .add_paragraph ()
        p .add_run (title )
        p .add_run (f"\t\t\t\t\t\t\t\t\t\t{page }").font .color .rgb =RGBColor (100 ,116 ,139 )
    doc .add_page_break ()




    add_h1 ("List of Figures & Tables")
    add_h2 ("List of Figures")
    doc .add_paragraph ("Figure 6.1: High-Level Three-Tier Architecture Diagram of Nexus Pathology\nFigure 7.1: Entity-Relationship (ER) Diagram of the Database Schema\nFigure 9.1: Machine Learning Lifecycle & Leakage Prevention Pipeline\nFigure 13.1: Synthetic Data Experiment Augmentation Workflow\nFigure 14.1: Security & IDOR Verification Defense Model")

    add_h2 ("List of Tables")
    doc .add_paragraph ("Table 4.1: Comprehensive Existing vs. Proposed System Comparison Matrix\nTable 5.1: Hardware & Software Requirements Specification\nTable 7.1: Data Dictionary for Database Tables\nTable 10.1: Clinical Feature Specifications for 5 Disease Modules\nTable 12.1: Final Validated Model Performance Benchmarks\nTable 13.1: Controlled Synthetic Augmentation Evaluation Results\nTable 15.1: Automated Verification & Security Test Suite Matrix")
    doc .add_page_break ()




    add_h1 ("1. Introduction")
    doc .add_paragraph (
    "Diagnostic clinical pathology is the backbone of patient care, contributing to clinical decision-making across outpatient and inpatient settings. As medical facilities transition toward digital health records, pathology laboratories require specialized platforms capable of handling structured biomarker parameters, maintaining strict audit trails, and enabling secure patient access.\n\n"
    "Simultaneously, machine learning has emerged as a promising tool for computational decision support, capable of recognizing subtle multivariate biomarker interactions. However, deploying machine learning in medical environments requires extreme care to prevent data leakage, preserve report integrity, and maintain clear boundaries between clinical facts and statistical predictions.\n\n"
    "Nexus Pathology was developed to address this intersection: providing a robust, web-based digital pathology management system with an integrated, strictly decoupled experimental machine-learning decision-support backend."
    )




    add_h1 ("2. Problem Statement")
    doc .add_paragraph (
    "Clinical laboratories face major operational, technological, and privacy challenges:\n"
    "1. Fragmented Record Keeping: Reliance on physical paper records or unstructured documents leads to delayed patient follow-up and loss of historical records.\n"
    "2. Insecure Patient Delivery: Transmitting medical reports via unencrypted emails or unauthenticated web links creates severe patient privacy vulnerabilities.\n"
    "3. Lack of Structured Biochemical Data: Without structured biomarker modeling, automated range validation and longitudinal trend analysis remain impossible.\n"
    "4. Risk of Diagnostic AI Overreach: Unregulated AI applications risk overwriting legal laboratory records or presenting statistical predictions as definitive clinical diagnoses without disclaimers."
    )




    add_h1 ("3. Project Objectives")
    doc .add_paragraph (
    "The project aims to achieve the following specific objectives:\n"
    "• Digitize clinical laboratory reporting with Draft and Finalized lifecycle workflows.\n"
    "• Implement an authenticated, IDOR-protected Patient Portal with cryptographic PBKDF2 hashing.\n"
    "• Strictly decouple official pathology reports from probabilistic ML inferences.\n"
    "• Train, validate, and serialize five specialized diagnostic ML pipelines (Anemia, Dengue, Liver, Thyroid, Malaria).\n"
    "• Empirically investigate synthetic tabular data augmentation (+25%, +50%, +100%).\n"
    "• Enforce comprehensive cybersecurity controls and achieve 100% automated test verification."
    )




    add_h1 ("4. Existing vs. Proposed System")
    tbl_comp =doc .add_table (rows =6 ,cols =3 )
    tbl_comp .alignment =WD_TABLE_ALIGNMENT .CENTER 
    headers =["Evaluation Dimension","Existing Conventional System","Proposed Nexus Pathology Platform"]
    for i ,h in enumerate (headers ):
        cell =tbl_comp .cell (0 ,i )
        set_cell_background (cell ,"E2E8F0")
        p =cell .paragraphs [0 ]
        p .add_run (h ).font .bold =True 

    rows_data =[
    ("Record Storage","Physical papers / flat spreadsheets","Relational SQLite DB with foreign key constraints"),
    ("Patient Access","In-person collection / unencrypted email","Authenticated Patient Portal with PBKDF2 PIN hashing"),
    ("Decision Support","None / disconnected external tools","5 integrated, decoupled ML decision-support pipelines"),
    ("Report Integrity","Casual edits without audit tracking","Official reports are strictly immutable by ML analysis"),
    ("Cybersecurity","Minimal access controls / IDOR vulnerabilities","RBAC, signed HMAC tokens, parameterized SQL, 100% test pass")
    ]
    for r_idx ,(dim ,ex ,pr )in enumerate (rows_data ,start =1 ):
        tbl_comp .cell (r_idx ,0 ).paragraphs [0 ].add_run (dim ).font .bold =True 
        tbl_comp .cell (r_idx ,1 ).paragraphs [0 ].add_run (ex )
        tbl_comp .cell (r_idx ,2 ).paragraphs [0 ].add_run (pr )

    doc .add_paragraph ().paragraph_format .space_after =Pt (6 )




    add_h1 ("5. System Architecture & Requirements")
    doc .add_paragraph (
    "Nexus Pathology employs a decoupled three-tier architecture comprising:\n"
    "1. Presentation Tier: Responsive Vanilla HTML5/CSS3/JavaScript SPA with dedicated print formatting (@media print).\n"
    "2. Application Tier: FastAPI asynchronous REST backend with Pydantic request validation and RBAC middleware.\n"
    "3. Persistence Tier: SQLite database (`pathology.db`) paired with serialized Scikit-Learn pipelines in `models/`."
    )
    add_h2 ("System Requirements")
    doc .add_paragraph (
    "• Hardware: Quad-core CPU, 4GB RAM minimum (8GB recommended), 2GB disk space.\n"
    "• Software: Python 3.10+, FastAPI 0.115+, Uvicorn 0.34+, Scikit-Learn 1.6+, OpenCV 4.10+, SQLite 3."
    )




    add_h1 ("6. Database Design & Decoupled Schema")
    doc .add_paragraph (
    "The relational schema contains four core tables designed to guarantee clinical data integrity:\n"
    "• `patients`: Stores demographic records and PBKDF2-hashed access PINs.\n"
    "• `lab_reports`: Stores authoritative official laboratory reports (patient ID, panel, parameters JSON, remarks, status).\n"
    "• `ml_predictions`: Stores an immutable audit trail of all decision support inferences (input snapshot, model version, confidence, disclaimer).\n"
    "• `users`: Stores laboratory staff credentials and administrative roles."
    )




    add_h1 ("7. Machine Learning Methodology & Results")
    doc .add_paragraph (
    "Five specialized pipelines were trained with zero data leakage (transformers fitted strictly on training folds) and evaluated via 5-fold cross-validation:"
    )

    tbl_ml =doc .add_table (rows =6 ,cols =5 )
    tbl_ml .alignment =WD_TABLE_ALIGNMENT .CENTER 
    ml_headers =["Disease / Panel","Algorithm","Features","Holdout Accuracy","5-Fold Cross Validation"]
    for i ,h in enumerate (ml_headers ):
        cell =tbl_ml .cell (0 ,i )
        set_cell_background (cell ,"E2E8F0")
        p =cell .paragraphs [0 ]
        p .add_run (h ).font .bold =True 

    ml_data =[
    ("Anemia (CBC)","Logistic Regression","11 Features","100.00%","95.49% ± 1.64%"),
    ("Dengue Hematology","Random Forest","8 Features","92.93%","91.30% ± 2.36%"),
    ("Liver Disease (LFT)","Gradient Boosting","10 Features","72.81%","69.30% ± 2.94% (95.06% Recall)"),
    ("Thyroid Profile","Multinomial Logistic Reg","5 Features","100.00%","95.81% ± 3.09%"),
    ("Malaria Microscopy","Gradient Boosting + CV","354 Features","94.03%","Strict Unseen (97.80% Recall)")
    ]
    for r_idx ,row in enumerate (ml_data ,start =1 ):
        for c_idx ,val in enumerate (row ):
            tbl_ml .cell (r_idx ,c_idx ).paragraphs [0 ].add_run (val )

    doc .add_paragraph ().paragraph_format .space_after =Pt (6 )
    doc .add_paragraph (
    "Malaria Dataset Deduplication: A cryptographic SHA-256 duplicate audit identified 25 duplicate images between train and test partitions. Purging duplicates and retraining ensured strict unseen generalization (94.03% accuracy, 97.80% recall, 95.70% F1)."
    )




    add_h1 ("8. Controlled Synthetic Data Experiment")
    doc .add_paragraph (
    "A controlled experiment evaluated whether synthetic tabular data augmentation (+25%, +50%, +100%) improved model generalization. Holdout test sets remained strictly quarantined, and synthesis was performed strictly in-fold.\n\n"
    "Findings: Real baseline data matched or outperformed synthetic augmentations. In Liver Disease, 100% synthetic data degraded holdout accuracy from 72.81% down to 66.67%. Consequently, synthetic data is strictly excluded from production."
    )




    add_h1 ("9. Cybersecurity & Verification Testing")
    doc .add_paragraph (
    "Security controls include PBKDF2-HMAC-SHA256 password/PIN hashing, signed HMAC session tokens, RBAC, strict IDOR prevention, parameterized SQL queries, 5MB image upload caps with OpenCV decoding, and automated online database backups.\n\n"
    "Automated Test Suite: 25 distinct test scenarios (15 security & IDOR tests, 10 integration and API tests) were executed with a 100% pass rate in 0.35 seconds."
    )




    add_h1 ("10. Conclusion & Future Scope")
    doc .add_paragraph (
    "Conclusion: Nexus Pathology successfully demonstrates a digital pathology management platform with decoupled, safe ML decision support, robust cybersecurity, and complete test verification.\n\n"
    "Future Scope: Integration of Explainable AI (SHAP/LIME), HL7/FHIR interoperability, PostgreSQL migration, multi-factor authentication, and formal clinical validation."
    )




    add_h1 ("11. References & Appendix")
    doc .add_paragraph (
    "References Placeholder:\n"
    "[1] Scikit-Learn: Machine Learning in Python, Pedregosa et al., JMLR 12, pp. 2825-2830, 2011.\n"
    "[2] FastAPI: Modern, Fast Web Framework for Python, Sebastián Ramírez, 2018.\n"
    "[3] SQLite Database Engine, D. Richard Hipp et al., 2000.\n"
    "[4] OpenCV: Open Source Computer Vision Library, Bradski, G., 2000.\n\n"
    "Appendix: Comprehensive Viva Questions and Project Presentation Outline are available in the project `docs/` repository."
    )

    os .makedirs ("report",exist_ok =True )
    out_docx =os .path .abspath ("report/Nexus_Pathology_Final_Project_Report.docx")
    doc .save (out_docx )
    print (f"Final Project Report saved successfully to: {out_docx }")

if __name__ =="__main__":
    create_report ()
